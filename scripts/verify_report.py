import docx
from docx.shared import Inches, Pt

doc = docx.Document(r'C:\Users\sahil\Documents\Ia-2\Digital_Twin_FM_IoT_Simulation_Telemetry_Report.docx')

print("--- DOCUMENT VERIFICATION METRICS ---")
print("Total Paragraphs:", len(doc.paragraphs))
print("Total Tables:", len(doc.tables))

word_count = sum(len(p.text.split()) for p in doc.paragraphs)
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            word_count += sum(len(p.text.split()) for p in cell.paragraphs)

print("Total Word Count:", word_count)

# Margins check
sec = doc.sections[0]
print("Top Margin (in):", sec.top_margin.inches)
print("Bottom Margin (in):", sec.bottom_margin.inches)
print("Left Margin (in):", sec.left_margin.inches)
print("Right Margin (in):", sec.right_margin.inches)
print("Page Width (mm):", sec.page_width.mm)
print("Page Height (mm):", sec.page_height.mm)

# Estimated pages (in Word, Times New Roman 12pt with 1.15 spacing and tables/figures yields ~250 words per page):
est_pages = word_count / 250.0
print(f"Estimated Formatted Page Count in Word: ~{est_pages:.1f} Pages")
