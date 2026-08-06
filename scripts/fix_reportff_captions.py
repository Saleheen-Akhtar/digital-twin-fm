import docx
import win32com.client
import os

def fix_doc_captions(file_path):
    doc = docx.Document(file_path)
    print(f"Loaded {file_path}")
    print(f"Total Paragraphs: {len(doc.paragraphs)}, Total Tables: {len(doc.tables)}")

    # Move table captions above tables
    for t_idx, t in enumerate(doc.tables):
        tbl_elem = t._element
        p_after = tbl_elem.getnext()
        
        if p_after is not None and p_after.tag.endswith('p'):
            # Convert CT_P element to docx Paragraph
            p_obj = docx.text.paragraph.Paragraph(p_after, doc)
            txt = p_obj.text.strip()
            
            if txt.startswith('Table'):
                if 'Table 8.1' in txt:
                    p_obj.text = txt.replace('Table 8.1', 'Table 5.1')
                    txt = p_obj.text.strip()
                
                # Move p_after element before tbl_elem
                tbl_elem.addprevious(p_after)
                print(f"Moved caption ABOVE Table {t_idx+1}: '{txt}'")

    # Correct figure caption numbers
    for p in doc.paragraphs:
        txt = p.text.strip()
        if 'Figure 4.8' in txt:
            p.text = txt.replace('Figure 4.8', 'Figure 4.3')
            print(f"Corrected caption: '{p.text}'")
        elif 'Figure 9.1' in txt:
            p.text = txt.replace('Figure 9.1', 'Figure 6.1')
            print(f"Corrected caption: '{p.text}'")

    out1 = r'C:\Users\sahil\Downloads\Saleheen_Major_project_reportff.docx'
    out2 = r'C:\Users\sahil\Downloads\Saleheen_Major_project_report.docx'
    out3 = r'C:\Users\sahil\Documents\Ia-2\Saleheen_Major_project_report.docx'

    doc.save(out1)
    doc.save(out2)
    doc.save(out3)
    print("Saved updated files to Downloads and Documents folders.")

    # Close MS Word first if running
    os.system("powershell -Command \"Stop-Process -Name WINWORD -Force -ErrorAction SilentlyContinue\"")

    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    wdoc = word.Documents.Open(out1)
    pages = wdoc.ComputeStatistics(2)
    wdoc.Close(False)
    word.Quit()
    print(f"=== UPDATED MS WORD RENDERED PAGE COUNT = {pages} PAGES ===")

if __name__ == '__main__':
    fix_doc_captions(r'C:\Users\sahil\Downloads\Saleheen_Major_project_reportff.docx')
