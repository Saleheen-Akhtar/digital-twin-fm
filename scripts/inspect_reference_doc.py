import docx

doc_path = r'C:\Users\sahil\Downloads\MCA_MajorProjectReport_ContentPages_N.docx'
doc = docx.Document(doc_path)

print('=== SECTION MARGINS & PAGE SETUP ===')
for idx, sec in enumerate(doc.sections):
    print(f'Section {idx+1}:')
    print(f'  Width: {sec.page_width.mm:.1f} mm, Height: {sec.page_height.mm:.1f} mm')
    print(f'  Top: {sec.top_margin.inches:.2f}\", Bottom: {sec.bottom_margin.inches:.2f}\"')
    print(f'  Left: {sec.left_margin.inches:.2f}\", Right: {sec.right_margin.inches:.2f}\"')

print('\n=== PARAGRAPHS & STYLES (Sample) ===')
for idx, p in enumerate(doc.paragraphs[:60]):
    text = p.text.strip()
    if text:
        fonts = set(r.font.name for r in p.runs if r.font.name)
        sizes = set(r.font.size.pt if r.font.size else None for r in p.runs)
        bolds = set(r.font.bold for r in p.runs if r.font.bold is not None)
        print(f'[{idx+1}] Style: {p.style.name} | Align: {p.alignment} | Fonts: {fonts} | Sizes: {sizes} | Bold: {bolds}')
        print(f'     Text: {text[:100]}')

print('\n=== TABLES IN TEMPLATE ===')
print(f'Total Tables: {len(doc.tables)}')
for idx, t in enumerate(doc.tables):
    print(f'Table {idx+1}: {len(t.rows)} rows, {len(t.columns)} cols, Style: {t.style.name}')

print('\n=== ALL HEADINGS & CHAPTER TITLES ===')
for p in doc.paragraphs:
    text = p.text.strip()
    if text and (p.style.name.startswith('Heading') or text.isupper() or 'CHAPTER' in text.upper() or 'CONTENTS' in text.upper()):
        print(f'[{p.style.name}] {text}')
